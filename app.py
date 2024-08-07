import streamlit as st
from sqlalchemy.orm import Session
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from models import Document, UserHistory, Base
from cryptography.fernet import Fernet
from PyPDF2 import PdfReader  # Updated import
from docx import Document as DocxDocument
import os

# Database setup
DATABASE_URL = "sqlite:///./test.db"
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base.metadata.create_all(bind=engine)

# Encryption key setup
KEY = Fernet.generate_key()
cipher_suite = Fernet(KEY)

def save_file(file):
    temp_filename = f"temp_{file.name}"
    with open(temp_filename, "wb") as f:
        f.write(file.getbuffer())
    return temp_filename

def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as f:
        reader = PdfReader(f)  # Updated class
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def add_document_to_db(file_path, filename, db: Session):
    if file_path.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    elif file_path.endswith('.txt'):
        with open(file_path, 'r') as file:
            text = file.read()
    else:
        return "Unsupported file format"
    
    encrypted_content = cipher_suite.encrypt(text.encode())
    db.add(Document(filename=filename, content=encrypted_content))
    db.commit()

def query_documents(query, db: Session):
    results = []
    for doc in db.query(Document).all():
        decrypted_content = cipher_suite.decrypt(doc.content).decode()
        if query.lower() in decrypted_content.lower():
            results.append({
                'filename': doc.filename,
                'content': decrypted_content
            })
    return results

def main():
    st.title("Document Query Application")

    uploaded_file = st.file_uploader("Upload a document", type=['pdf', 'docx', 'txt'])
    if uploaded_file:
        file_path = save_file(uploaded_file)
        db = SessionLocal()
        add_document_to_db(file_path, uploaded_file.name, db)
        st.success("Document added successfully")

    query = st.text_input("Enter your query")
    if query:
        db = SessionLocal()
        results = query_documents(query, db)
        if results:
            for result in results:
                st.subheader(result['filename'])
                st.write(result['content'])
        else:
            st.write("No results found")

    st.subheader("Download Chat History")
    # Implement functionality to download chat history

if __name__ == "__main__":
    main()
